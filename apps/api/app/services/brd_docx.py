from __future__ import annotations

from io import BytesIO
import re
import unicodedata
from typing import Any
from urllib.parse import quote

from docx import Document


def build_brd_docx(title: str, markdown: str) -> bytes:
    document = Document()
    document.core_properties.title = title
    blocks = parse_markdown(markdown)

    if not blocks:
        document.add_heading(title, level=0)
    else:
        for block in blocks:
            render_block(document, block)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def sanitize_docx_filename(title: str) -> str:
    sanitized = re.sub(r'[\\/:*?"<>|]+', "-", title).strip()
    return sanitized or "BRD"


def build_docx_content_disposition(filename: str) -> str:
    sanitized = sanitize_docx_filename(filename)
    ascii_fallback = unicodedata.normalize("NFKD", sanitized).encode("ascii", "ignore").decode("ascii")
    ascii_fallback = re.sub(r"[^A-Za-z0-9._ -]+", "", ascii_fallback).strip() or "BRD"
    utf8_filename = quote(sanitized.encode("utf-8"))
    return f"""attachment; filename="{ascii_fallback}"; filename*=UTF-8''{utf8_filename}"""


def render_block(document: Document, block: dict[str, Any]) -> None:
    block_type = block["type"]

    if block_type == "heading":
        level = min(max(block["level"], 0), 4)
        document.add_heading(block["text"], level=level)
        return

    if block_type == "paragraph":
        document.add_paragraph(block["text"])
        return

    if block_type == "unordered-list":
        for item in block["items"]:
            document.add_paragraph(item["text"], style="List Bullet")
            for child in item["children"]:
                document.add_paragraph(child, style="List Bullet 2")
        return

    if block_type == "ordered-list":
        for item in block["items"]:
            document.add_paragraph(item["text"], style="List Number")
            for child in item["children"]:
                document.add_paragraph(child, style="List Bullet 2")
        return

    if block_type == "table":
        render_table(document, block["table"])
        return

    if block_type == "figure":
        figure = block["figure"]
        placeholder = document.add_paragraph()
        placeholder.add_run(figure["alt"]).bold = True
        placeholder.add_run(f" ({figure['src']})")
        if figure["caption"]:
            caption = document.add_paragraph(figure["caption"])
            caption.style = "Intense Quote"
        return

    if block_type == "rule":
        document.add_paragraph("--------------------------------------------------")


def render_table(document: Document, table_data: dict[str, Any]) -> None:
    headers: list[str] = table_data["headers"]
    rows: list[list[str]] = table_data["rows"]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = header

    for row in rows:
        cells = table.add_row().cells
        for index, cell in enumerate(row):
            cells[index].text = cell


def parse_markdown(markdown: str) -> list[dict[str, Any]]:
    lines = markdown.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    blocks: list[dict[str, Any]] = []
    index = 0

    while index < len(lines):
        trimmed = lines[index].strip()
        if not trimmed:
            index += 1
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.+)$", trimmed)
        if heading_match:
            blocks.append(
                {
                    "type": "heading",
                    "level": len(heading_match.group(1)),
                    "text": heading_match.group(2).strip(),
                }
            )
            index += 1
            continue

        if re.match(r"^---+$", trimmed):
            blocks.append({"type": "rule"})
            index += 1
            continue

        figure_match = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)$", trimmed)
        if figure_match:
            next_trimmed = lines[index + 1].strip() if index + 1 < len(lines) else ""
            caption = (
                next_trimmed
                if re.match(r"^(Hình|Figure)\s+\d+:", next_trimmed, flags=re.IGNORECASE)
                else None
            )
            blocks.append(
                {
                    "type": "figure",
                    "figure": {
                        "alt": figure_match.group(1).strip() or "Hình minh họa",
                        "src": figure_match.group(2).strip(),
                        "caption": caption,
                    },
                }
            )
            index += 2 if caption else 1
            continue

        next_trimmed = lines[index + 1].strip() if index + 1 < len(lines) else ""
        if is_table_header(trimmed, next_trimmed):
            table, index = parse_table(lines, index)
            blocks.append({"type": "table", "table": table})
            continue

        if unordered_list_match(trimmed):
            items, index = parse_list(lines, index, ordered=False)
            blocks.append({"type": "unordered-list", "items": items})
            continue

        if ordered_list_match(trimmed):
            items, index = parse_list(lines, index, ordered=True)
            blocks.append({"type": "ordered-list", "items": items})
            continue

        paragraph_lines = [trimmed]
        index += 1
        while index < len(lines):
            next_trimmed = lines[index].strip()
            if (
                not next_trimmed
                or re.match(r"^(#{1,6})\s+(.+)$", next_trimmed)
                or re.match(r"^---+$", next_trimmed)
                or re.match(r"^!\[([^\]]*)\]\(([^)]+)\)$", next_trimmed)
                or is_table_header(next_trimmed, lines[index + 1].strip() if index + 1 < len(lines) else "")
                or unordered_list_match(next_trimmed)
                or ordered_list_match(next_trimmed)
            ):
                break
            paragraph_lines.append(next_trimmed)
            index += 1
        blocks.append({"type": "paragraph", "text": " ".join(paragraph_lines)})

    return blocks


def parse_list(lines: list[str], start_index: int, ordered: bool) -> tuple[list[dict[str, Any]], int]:
    items: list[dict[str, Any]] = []
    index = start_index

    while index < len(lines):
        trimmed = lines[index].strip()
        if not trimmed:
            break

        is_match = ordered_list_match(trimmed) if ordered else unordered_list_match(trimmed)
        if not is_match:
            break

        text = re.sub(r"^\d+\.\s+" if ordered else r"^-\s+", "", trimmed).strip()
        item = {"text": text, "children": []}
        index += 1

        while index < len(lines):
            child_line = lines[index]
            if not child_line.strip():
                index += 1
                break
            if re.match(r"^\s+-\s+", child_line):
                item["children"].append(re.sub(r"^\s+-\s+", "", child_line).strip())
                index += 1
                continue
            if re.match(r"^\s+", child_line) and not item["children"]:
                item["text"] = f"{item['text']} {child_line.strip()}".strip()
                index += 1
                continue
            break

        items.append(item)

    return items, index


def parse_table(lines: list[str], start_index: int) -> tuple[dict[str, Any], int]:
    headers = parse_table_cells(lines[start_index].strip())
    rows: list[list[str]] = []
    index = start_index + 2

    while index < len(lines):
        trimmed = lines[index].strip()
        if not trimmed or "|" not in trimmed:
            break
        rows.append(parse_table_cells(trimmed))
        index += 1

    return {"headers": headers, "rows": rows}, index


def parse_table_cells(line: str) -> list[str]:
    parts = [part.strip() for part in line.split("|")]
    return [
        part
        for index, part in enumerate(parts)
        if not (index == 0 and part == "") and not (index == len(parts) - 1 and part == "")
    ]


def is_table_header(line: str, next_line: str) -> bool:
    normalized_line = line.strip()
    normalized_next_line = next_line.strip()
    return (
        normalized_line.startswith("|")
        and normalized_line.endswith("|")
        and normalized_next_line.startswith("|")
        and normalized_next_line.endswith("|")
        and "---" in normalized_next_line
    )


def unordered_list_match(value: str) -> bool:
    return bool(re.match(r"^-\s+", value))


def ordered_list_match(value: str) -> bool:
    return bool(re.match(r"^\d+\.\s+", value))
